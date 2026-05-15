"""
Pricing Calculator Module

Handles the calculation of composite rates including markup and tax
for vendor quotes to generate client-facing pricing.

Outputs:
  - Client sheet: Description | Composite Unit Rate | Total (3 cols, all-in price)
  - Audit sheet: Full formula trail for leadership verification
"""

import pandas as pd
import numpy as np
import re
import os
import tempfile
from typing import Optional, Dict, Any, Tuple, List
from io import BytesIO


class PricingProcessor:
    """
    Processes vendor quotes and calculates composite rates with markup and tax.
    
    Attributes:
        tax_rate (float): Tax rate as decimal (default: 0.0825 for 8.25%)
        margin_rate (float): Margin/markup rate as decimal (default: 0.10 for 10%)
        shipping_cost (float): Pass-through shipping cost (default: 0.00)
    """
    
    def __init__(self, tax_rate: float = 0.0825, margin_rate: float = 0.10,
                 shipping_cost: float = 0.0):
        self.tax_rate = tax_rate
        self.margin_rate = margin_rate
        self.shipping_cost = shipping_cost
    
    def calculate_composite_rate(self, unit_cost: float) -> float:
        """
        Calculate the composite unit rate: unit_cost × (1 + margin) × (1 + tax).
        All-in price per unit — margin first, then tax baked in.
        
        Returns value with 6 decimal precision to prevent rounding errors
        on high-quantity items (e.g., 44,000 ft @ $0.829917/ft).
        """
        if pd.isna(unit_cost) or unit_cost == 0:
            return 0.0
        composite_rate = unit_cost * (1 + self.margin_rate) * (1 + self.tax_rate)
        return round(composite_rate, 6)
    
    # ------------------------------------------------------------------
    # Data cleaning
    # ------------------------------------------------------------------
    def clean_vendor_quote(self, df: pd.DataFrame,
                           desc_col: str, qty_col: str,
                           cost_col: str) -> Tuple[pd.DataFrame, float, float]:
        """
        Clean a raw parsed vendor quote DataFrame.
        
        Phase 1 — Merge multi-line descriptions: vendor PDFs often wrap a
                   single line-item description across 2-4 rows.  Continuation
                   rows (no qty, no cost) are appended to the previous real row.
        Phase 2 — Extract sales tax and freight dollar amounts.
        Phase 3 — Strip junk: vendor headers, repeated column-header rows,
                   empty rows, subtotal markers.
        
        Returns:
            (cleaned_df, sales_tax_dollars, freight_dollars)
        """
        # ── Phase 1: merge continuation rows ──────────────────────
        df = self._merge_multiline_rows(df, desc_col, qty_col, cost_col)
        
        # ── Phase 1b: strip commas from qty/cost so "25,035" → 25035 ──
        for col in [qty_col, cost_col]:
            df[col] = df[col].astype(str).str.replace(',', '', regex=False)
            df[col] = pd.to_numeric(df[col], errors='coerce')
        
        # ── Phase 2 & 3: extract tax/freight, strip junk ─────────
        sales_tax = 0.0
        freight = 0.0
        
        header_pattern = re.compile(
            r'^(item|description|u/?m|total|qty|quantity|rate|none|tbd|'
            r'net\s*\d+|a\.?b\.?|estimate|terms|rep|project|col\d+|page)$',
            re.IGNORECASE
        )
        tax_pattern = re.compile(r'sales\s*tax|^tax$', re.IGNORECASE)
        freight_pattern = re.compile(r'freight|shipping', re.IGNORECASE)
        subtotal_pattern = re.compile(r'subtotal|sub\s*total|grand\s*total', re.IGNORECASE)
        
        keep_mask = []
        
        for _, row in df.iterrows():
            desc_val = str(row.get(desc_col, '')).strip()
            qty_val = pd.to_numeric(row.get(qty_col), errors='coerce')
            cost_val = pd.to_numeric(row.get(cost_col), errors='coerce')
            # Build row_text from ALL columns for keyword searching
            row_text = ' '.join(str(v) for v in row.values if pd.notna(v))
            
            # --- Extract sales tax ---
            if tax_pattern.search(row_text):
                # Search ALL cells for the dollar amount (could be in any column)
                for v in row.values:
                    parsed = _parse_currency(v)
                    if parsed and parsed > sales_tax:
                        sales_tax = parsed
                keep_mask.append(False)
                continue
            
            # --- Extract freight ---
            if freight_pattern.search(row_text):
                # Look for a real dollar amount > 0; ignore qty=1 / cost=0 pattern
                row_freight = 0.0
                for v in row.values:
                    parsed = _parse_currency(v)
                    if parsed and parsed > row_freight:
                        row_freight = parsed
                if row_freight > freight:
                    freight = row_freight
                keep_mask.append(False)
                continue
            
            # --- Remove subtotal rows ---
            if subtotal_pattern.search(row_text):
                keep_mask.append(False)
                continue
            
            # --- Remove repeated column-header / junk rows ---
            if not desc_val or desc_val.lower() == 'none' or header_pattern.match(desc_val):
                keep_mask.append(False)
                continue
            
            # --- Remove rows where qty AND cost are both zero or NaN ---
            if (pd.isna(qty_val) or qty_val == 0) and (pd.isna(cost_val) or cost_val == 0):
                keep_mask.append(False)
                continue
            
            keep_mask.append(True)
        
        cleaned_df = df.loc[df.index[keep_mask]].copy().reset_index(drop=True)
        
        # Ensure numeric columns (already stripped commas above, but be safe)
        cleaned_df[qty_col] = pd.to_numeric(cleaned_df[qty_col], errors='coerce').fillna(0)
        cleaned_df[cost_col] = pd.to_numeric(cleaned_df[cost_col], errors='coerce').fillna(0)
        
        return cleaned_df, sales_tax, freight
    
    @staticmethod
    def _merge_multiline_rows(df: pd.DataFrame, desc_col: str,
                               qty_col: str, cost_col: str) -> pd.DataFrame:
        """
        Merge multi-line PDF rows into single line-items.
        
        Vendor PDFs typically produce this pattern for each item:
          Row A: part# + description text  (no qty, no cost)
          Row B: continuation description  (no qty, no cost)
          Row C: *** IN-STOCK note ***     (no qty, no cost)
          Row D: 150  ea  8.20  1,230.00T  (HAS qty and cost — the data row)
        
        Strategy: always buffer description-only rows.  When a data row
        arrives (has qty or cost), prepend all buffered descriptions onto
        it, clear the buffer, and emit the merged row.
        """
        # Pattern for junk lines that should never be part of a description.
        # Also matches short generic strings (<=3 words, no part-number chars)
        # that typically appear in vendor page headers.
        junk_line = re.compile(
            r'^(date|estimate|tbd|net\s*\d+|a\.?b\.?|item|description|'
            r'u/?m|total|qty|quantity|rate|terms|rep|project|page|'
            r'col\d+|none|nan)$',
            re.IGNORECASE
        )
        # Short lines (1-3 words, no digits or hyphens) are likely vendor
        # header labels like "Gotebo Resound", "A.B.", "Rep", etc.
        short_header = re.compile(r'^[A-Za-z\s.]+$')
        
        merged_rows: List[pd.Series] = []
        desc_buffer: List[str] = []
        
        for _, row in df.iterrows():
            qty_val = pd.to_numeric(row.get(qty_col), errors='coerce')
            cost_val = pd.to_numeric(row.get(cost_col), errors='coerce')
            desc_val = str(row.get(desc_col, '')).strip()
            
            has_qty = pd.notna(qty_val) and qty_val != 0
            has_cost = pd.notna(cost_val) and cost_val != 0
            
            if has_qty or has_cost:
                # ── Data row ──
                data_row = row.copy()
                
                # Filter junk lines out of the buffer before prepending
                clean_buffer = [
                    ln for ln in desc_buffer
                    if not junk_line.match(ln.strip())
                    and not (short_header.match(ln.strip()) and len(ln.split()) <= 3)
                ]
                
                if clean_buffer:
                    combined_desc = ' '.join(clean_buffer)
                    existing = desc_val
                    # If the data row's desc cell is just a number (the qty
                    # that bled into the desc column), replace it entirely
                    try:
                        float(existing.replace(',', ''))
                        data_row[desc_col] = combined_desc
                    except ValueError:
                        if existing:
                            data_row[desc_col] = (combined_desc + ' ' + existing).strip()
                        else:
                            data_row[desc_col] = combined_desc
                desc_buffer = []
                
                merged_rows.append(data_row)
            else:
                # ── Description-only / note row — always buffer ──
                if desc_val and desc_val.lower() not in ('none', 'nan', ''):
                    desc_buffer.append(desc_val)
        
        if not merged_rows:
            return df
        
        return pd.DataFrame(merged_rows).reset_index(drop=True)
    
    # ------------------------------------------------------------------
    # Quote processing
    # ------------------------------------------------------------------
    def process_quote(self, df: pd.DataFrame, 
                      desc_col: str = None, qty_col: str = None, 
                      cost_col: str = None) -> pd.DataFrame:
        """
        Process a cleaned vendor quote DataFrame and add calculated columns.
        
        Adds:
          - Composite Unit Rate = unit_cost × (1 + margin) × (1 + tax)
          - Total Price = Composite Unit Rate × Quantity
        """
        processed_df = df.copy()
        
        processed_df[cost_col] = pd.to_numeric(processed_df[cost_col], errors='coerce').fillna(0)
        processed_df[qty_col] = pd.to_numeric(processed_df[qty_col], errors='coerce').fillna(0)
        
        processed_df['Composite Unit Rate'] = processed_df[cost_col].apply(
            self.calculate_composite_rate
        )
        
        processed_df['Total Price'] = (
            processed_df[qty_col] * processed_df['Composite Unit Rate']
        ).round(2)
        
        # Store mapped column names for downstream methods
        processed_df.attrs['desc_col'] = desc_col
        processed_df.attrs['qty_col'] = qty_col
        processed_df.attrs['cost_col'] = cost_col
        
        return processed_df
    
    # ------------------------------------------------------------------
    # Internal Audit Spreadsheet (full formula trail for leadership)
    # ------------------------------------------------------------------
    def generate_internal_audit_spreadsheet(self, df: pd.DataFrame, 
                                       tax_rate: float, 
                                       margin_rate: float,
                                       output_path=None,
                                       desc_col: str = None,
                                       qty_col: str = None,
                                       cost_col: str = None,
                                       shipping_cost: float = 0.0) -> pd.DataFrame:
        """
        Full audit trail: Vendor Cost → +Margin → After Margin → +Tax →
        Composite Unit Rate → Line Total → Formula.
        Shipping added as a separate line at the bottom.
        """
        if desc_col is None:
            desc_col = df.attrs.get('desc_col')
        if qty_col is None:
            qty_col = df.attrs.get('qty_col')
        if cost_col is None:
            cost_col = df.attrs.get('cost_col')
        
        margin_col = f'+Margin ({margin_rate*100:.1f}%)'
        tax_col = f'+Tax ({tax_rate*100:.2f}%)'
        
        audit_data = []
        for _, row in df.iterrows():
            if not (desc_col and qty_col):
                continue
            if pd.isna(row.get(desc_col)) or pd.isna(row.get(qty_col)):
                continue
            
            quantity = float(pd.to_numeric(row[qty_col], errors='coerce') or 0)
            unit_cost = float(pd.to_numeric(row.get(cost_col, 0), errors='coerce') or 0) if cost_col else 0.0

            # Use process_quote() outputs as the single source of truth so
            # audit and client exports always match at line level and totals.
            composite_rate = float(pd.to_numeric(row.get('Composite Unit Rate', 0), errors='coerce') or 0)
            line_total = float(pd.to_numeric(row.get('Total Price', 0), errors='coerce') or 0)

            # Keep audit breakdown columns while staying aligned with the
            # authoritative composite value above.
            if tax_rate != -1:
                after_margin = round(composite_rate / (1 + tax_rate), 6)
            else:
                after_margin = round(unit_cost * (1 + margin_rate), 6)

            margin_per_unit = round(after_margin - unit_cost, 6)
            tax_per_unit = round(composite_rate - after_margin, 6)
            row_num = len(audit_data) + 2  # Excel row (1-indexed, +1 for header)
            
            audit_data.append({
                'Description': row[desc_col],
                'Qty': quantity,
                'Vendor Unit Cost': unit_cost,
                margin_col: margin_per_unit,
                'After Margin': after_margin,
                tax_col: tax_per_unit,
                'Composite Unit Rate': composite_rate,
                'Line Total': line_total,
                'Formula': f'=C{row_num}*(1+{margin_rate})*(1+{tax_rate})'
            })
        
        audit_df = pd.DataFrame(audit_data)
        
        if audit_df.empty:
            return audit_df
        
        # Add shipping line if applicable
        if shipping_cost > 0:
            ship_row = {col: '' for col in audit_df.columns}
            ship_row['Description'] = 'Shipping / Freight (Pass-Through)'
            ship_row['Qty'] = 1
            ship_row['Vendor Unit Cost'] = shipping_cost
            ship_row[margin_col] = 0
            ship_row['After Margin'] = shipping_cost
            ship_row[tax_col] = 0
            ship_row['Composite Unit Rate'] = shipping_cost
            ship_row['Line Total'] = shipping_cost
            ship_row['Formula'] = 'Pass-through (no markup)'
            audit_df = pd.concat([audit_df, pd.DataFrame([ship_row])], ignore_index=True)
        
        # Totals row
        totals = {col: '' for col in audit_df.columns}
        totals['Description'] = 'TOTALS'
        for c in ['Qty', 'Vendor Unit Cost', margin_col, 'After Margin', tax_col,
                   'Composite Unit Rate', 'Line Total']:
            numeric_vals = pd.to_numeric(audit_df[c], errors='coerce')
            if c in ['Vendor Unit Cost', 'Composite Unit Rate']:
                totals[c] = ''
            else:
                totals[c] = round(numeric_vals.sum(), 2)
        totals['Line Total'] = round(pd.to_numeric(audit_df['Line Total'], errors='coerce').sum(), 2)
        totals['Formula'] = ''
        audit_df = pd.concat([audit_df, pd.DataFrame([totals])], ignore_index=True)
        
        # Write to Excel if output_path provided
        if output_path:
            self._write_audit_excel(audit_df, output_path, margin_col, tax_col)
        
        return audit_df
    
    def _write_audit_excel(self, audit_df, output_path, margin_col, tax_col):
        """Write the audit DataFrame to a formatted Excel file or BytesIO buffer."""
        is_buffer = hasattr(output_path, 'write')
        
        if is_buffer:
            with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
                temp_path = tmp.name
            target = temp_path
        else:
            target = output_path
        
        with pd.ExcelWriter(target, engine='xlsxwriter') as writer:
            audit_df.to_excel(writer, sheet_name='Audit_Details', index=False)
            workbook = writer.book
            worksheet = writer.sheets['Audit_Details']
            
            header_fmt = workbook.add_format({
                'bold': True, 'text_wrap': True, 'valign': 'top',
                'fg_color': '#D7E4BC', 'border': 1
            })
            money_fmt = workbook.add_format({'num_format': '$#,##0.00'})
            formula_fmt = workbook.add_format({'font_color': 'blue', 'italic': True})
            totals_fmt = workbook.add_format({
                'bold': True, 'bg_color': '#E2EFDA', 'border': 1
            })
            
            money_cols = ['Vendor Unit Cost', margin_col, 'After Margin',
                          tax_col, 'Composite Unit Rate', 'Line Total']
            
            for col_num, column in enumerate(audit_df.columns):
                worksheet.write(0, col_num, column, header_fmt)
                if column == 'Description':
                    worksheet.set_column(col_num, col_num, 45)
                elif column == 'Formula':
                    worksheet.set_column(col_num, col_num, 30)
                elif column in money_cols:
                    worksheet.set_column(col_num, col_num, 16, money_fmt)
                else:
                    worksheet.set_column(col_num, col_num, 12)
            
            # Formula column formatting
            if 'Formula' in audit_df.columns:
                fc = audit_df.columns.get_loc('Formula')
                for r in range(1, len(audit_df)):
                    val = audit_df.iloc[r]['Formula']
                    if val:
                        worksheet.write(r, fc, val, formula_fmt)
            
            # Totals row formatting
            tr = len(audit_df)  # xlsxwriter is 0-indexed but we wrote header at 0
            for col_num in range(len(audit_df.columns)):
                worksheet.write(len(audit_df), col_num,
                                audit_df.iloc[-1, col_num], totals_fmt)
        
        if is_buffer:
            with open(temp_path, 'rb') as f:
                output_path.write(f.read())
            try:
                os.unlink(temp_path)
            except OSError:
                pass
    
    # ------------------------------------------------------------------
    # Client Spreadsheet (3 columns: Description | Composite Unit Rate | Total)
    # ------------------------------------------------------------------
    def generate_client_spreadsheet(self, df: pd.DataFrame, 
                                   output_path=None,
                                   desc_col: str = None,
                                   qty_col: str = None,
                                   shipping_cost: float = 0.0) -> pd.DataFrame:
        """
        Generate a clean 3-column client-facing spreadsheet.
        Client sees: Description | Composite Unit Rate | Total
        No tax line, no margin line — one all-in price per unit.
        Shipping added as a flat pass-through line if > 0.
        """
        if desc_col is None:
            desc_col = df.attrs.get('desc_col')
        if qty_col is None:
            qty_col = df.attrs.get('qty_col')
        
        # Build the 3-column client view
        rows = []
        for _, row in df.iterrows():
            if desc_col and pd.notna(row.get(desc_col)):
                rows.append({
                    'Description': row[desc_col],
                    'Composite Unit Rate': row.get('Composite Unit Rate', 0),
                    'Total': row.get('Total Price', 0)
                })
        
        client_df = pd.DataFrame(rows)
        
        # Add shipping line
        if shipping_cost > 0:
            client_df = pd.concat([client_df, pd.DataFrame([{
                'Description': 'Shipping / Freight',
                'Composite Unit Rate': shipping_cost,
                'Total': shipping_cost
            }])], ignore_index=True)
        
        # Add totals row
        if not client_df.empty:
            total_sum = round(pd.to_numeric(client_df['Total'], errors='coerce').sum(), 2)
            totals_row = pd.DataFrame([{
                'Description': 'TOTAL',
                'Composite Unit Rate': '',
                'Total': total_sum
            }])
            client_df = pd.concat([client_df, totals_row], ignore_index=True)
        
        # Write to Excel if output_path provided
        if output_path:
            self._write_client_excel(client_df, output_path)
        
        return client_df
    
    def _write_client_excel(self, client_df, output_path):
        """Write the client DataFrame to a formatted Excel file or BytesIO buffer."""
        is_buffer = hasattr(output_path, 'write')
        
        if is_buffer:
            with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
                temp_path = tmp.name
            target = temp_path
        else:
            target = output_path
        
        with pd.ExcelWriter(target, engine='xlsxwriter') as writer:
            client_df.to_excel(writer, sheet_name='Client_Quote', index=False)
            workbook = writer.book
            worksheet = writer.sheets['Client_Quote']
            
            money_fmt = workbook.add_format({'num_format': '$#,##0.00'})
            header_fmt = workbook.add_format({
                'bold': True, 'fg_color': '#4472C4', 'font_color': 'white',
                'border': 1, 'text_wrap': True
            })
            totals_fmt = workbook.add_format({
                'bold': True, 'num_format': '$#,##0.00',
                'bg_color': '#D9E2F3', 'border': 1
            })
            totals_label_fmt = workbook.add_format({
                'bold': True, 'bg_color': '#D9E2F3', 'border': 1
            })
            
            for col_num, column in enumerate(client_df.columns):
                worksheet.write(0, col_num, column, header_fmt)
                if column == 'Description':
                    worksheet.set_column(col_num, col_num, 55)
                else:
                    worksheet.set_column(col_num, col_num, 22, money_fmt)
            
            # Format the totals row (last row)
            last_row = len(client_df)
            for col_num, column in enumerate(client_df.columns):
                if column == 'Description':
                    worksheet.write(last_row, col_num,
                                    client_df.iloc[-1][column], totals_label_fmt)
                elif column == 'Total':
                    worksheet.write(last_row, col_num,
                                    client_df.iloc[-1][column], totals_fmt)
                else:
                    worksheet.write(last_row, col_num,
                                    client_df.iloc[-1][column], totals_label_fmt)
        
        if is_buffer:
            with open(temp_path, 'rb') as f:
                output_path.write(f.read())
            try:
                os.unlink(temp_path)
            except OSError:
                pass
    
    # ------------------------------------------------------------------
    # Summary Report Generator (Markdown / .docx / Excel)
    # ------------------------------------------------------------------
    def generate_summary_report(self, processed_df: pd.DataFrame,
                                fmt: str = 'md',
                                shipping_cost: float = 0.0,
                                extracted_tax: float = 0.0,
                                extracted_freight: float = 0.0) -> str:
        """
        Generate a summary report of the processed quote.
        
        Args:
            processed_df: DataFrame after process_quote()
            fmt: Output format — 'md' for Markdown, 'docx' for Word
            shipping_cost: Pass-through shipping entered by user
            extracted_tax: Sales tax dollar amount extracted from vendor quote
            extracted_freight: Freight dollar amount extracted from vendor quote
            
        Returns:
            For 'md': Markdown string
            For 'docx': bytes of the .docx file
        """
        desc_col = processed_df.attrs.get('desc_col', 'Description')
        qty_col = processed_df.attrs.get('qty_col', 'Qty')
        cost_col = processed_df.attrs.get('cost_col', 'Unit Cost')
        
        # Calculate totals
        line_items = len(processed_df)
        vendor_subtotal = round(
            (pd.to_numeric(processed_df[cost_col], errors='coerce').fillna(0) *
             pd.to_numeric(processed_df[qty_col], errors='coerce').fillna(0)).sum(), 2)
        total_composite = round(processed_df['Total Price'].sum(), 2)
        total_margin = round(total_composite - vendor_subtotal - extracted_tax, 2)
        effective_margin_pct = round((total_margin / vendor_subtotal * 100), 2) if vendor_subtotal else 0
        grand_total = round(total_composite + shipping_cost, 2)
        
        if fmt == 'md':
            return self._summary_as_markdown(
                processed_df, desc_col, qty_col, line_items,
                vendor_subtotal, total_margin, effective_margin_pct,
                extracted_tax, extracted_freight, shipping_cost,
                total_composite, grand_total)
        elif fmt == 'docx':
            return self._summary_as_docx(
                processed_df, desc_col, qty_col, line_items,
                vendor_subtotal, total_margin, effective_margin_pct,
                extracted_tax, extracted_freight, shipping_cost,
                total_composite, grand_total)
        else:
            raise ValueError(f"Unsupported format: {fmt}")
    
    def _summary_as_markdown(self, df, desc_col, qty_col, line_items,
                              vendor_subtotal, total_margin, margin_pct,
                              extracted_tax, extracted_freight, shipping,
                              total_composite, grand_total) -> str:
        lines = [
            "# Quote Summary Report",
            "",
            f"**Line Items:** {line_items}",
            f"**Margin Rate:** {self.margin_rate*100:.1f}%",
            f"**Tax Rate:** {self.tax_rate*100:.2f}%",
            "",
            "## Line Items",
            "",
            "| Description | Composite Unit Rate | Total |",
            "|---|---|---|",
        ]
        for _, row in df.iterrows():
            desc = str(row.get(desc_col, ''))[:60]
            rate = f"${row.get('Composite Unit Rate', 0):,.2f}"
            total = f"${row.get('Total Price', 0):,.2f}"
            lines.append(f"| {desc} | {rate} | {total} |")
        
        lines.extend([
            "",
            "## Totals",
            "",
            f"| Vendor Subtotal | ${vendor_subtotal:,.2f} |",
            f"| Margin ({self.margin_rate*100:.1f}%) | ${total_margin:,.2f} ({margin_pct:.1f}% effective) |",
            f"| Extracted Sales Tax | ${extracted_tax:,.2f} |",
            f"| Extracted Freight | ${extracted_freight:,.2f} |",
            f"| Quote Total (with markup + tax) | ${total_composite:,.2f} |",
            f"| Pass-Through Shipping | ${shipping:,.2f} |",
            f"| **Grand Total** | **${grand_total:,.2f}** |",
        ])
        return "\n".join(lines)
    
    def _summary_as_docx(self, df, desc_col, qty_col, line_items,
                          vendor_subtotal, total_margin, margin_pct,
                          extracted_tax, extracted_freight, shipping,
                          total_composite, grand_total) -> bytes:
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx is required for .docx export. "
                              "Install it with: pip install python-docx")
        
        doc = Document()
        doc.add_heading('Quote Summary Report', 0)
        
        doc.add_paragraph(f'Line Items: {line_items}')
        doc.add_paragraph(f'Margin Rate: {self.margin_rate*100:.1f}%')
        doc.add_paragraph(f'Tax Rate: {self.tax_rate*100:.2f}%')
        
        doc.add_heading('Line Items', level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = 'Description'
        hdr[1].text = 'Composite Unit Rate'
        hdr[2].text = 'Total'
        
        for _, row in df.iterrows():
            r = table.add_row().cells
            r[0].text = str(row.get(desc_col, ''))[:80]
            r[1].text = f"${row.get('Composite Unit Rate', 0):,.2f}"
            r[2].text = f"${row.get('Total Price', 0):,.2f}"
        
        doc.add_heading('Totals', level=1)
        totals_data = [
            ('Vendor Subtotal', f'${vendor_subtotal:,.2f}'),
            (f'Margin ({self.margin_rate*100:.1f}%)',
             f'${total_margin:,.2f} ({margin_pct:.1f}% effective)'),
            ('Extracted Sales Tax', f'${extracted_tax:,.2f}'),
            ('Quote Total', f'${total_composite:,.2f}'),
            ('Pass-Through Shipping', f'${shipping:,.2f}'),
            ('Grand Total', f'${grand_total:,.2f}'),
        ]
        t2 = doc.add_table(rows=len(totals_data), cols=2)
        t2.style = 'Light Grid Accent 1'
        for i, (label, val) in enumerate(totals_data):
            t2.rows[i].cells[0].text = label
            t2.rows[i].cells[1].text = val
        
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()


# ======================================================================
# Module-level helpers
# ======================================================================

def _parse_currency(value) -> Optional[float]:
    """Try to parse a value as a dollar amount. Returns float or None."""
    if value is None or (isinstance(value, float) and np.isnan(value)):
        return None
    s = str(value).strip().replace('$', '').replace(',', '').replace('T', '')
    try:
        v = float(s)
        return v if v > 0 else None
    except (ValueError, TypeError):
        return None


def parse_hammon_spreadsheet(file_path: str, sheet_name: str = 'HAMMON OSP_UG_TAK') -> pd.DataFrame:
    """Parse the HAMMON spreadsheet format."""
    df = pd.read_excel(file_path, sheet_name=sheet_name, header=0)
    df = df.dropna(how='all')
    df = df[~df.iloc[:, 0].astype(str).str.contains('Subtotal|tax|Vendor Invoice|markup', na=False)]
    return df